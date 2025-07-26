from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from .models import *
import uuid
from uuid import uuid4
from django.contrib.auth import get_user_model
from django.http import JsonResponse
from django.db import transaction
from django.db.models import Count, Prefetch
from .forms import *
from django.forms import modelformset_factory
from django.contrib import messages
from collections import OrderedDict
from django.conf import settings
import requests
from django.contrib.auth import authenticate, login, logout
from django.utils.http import url_has_allowed_host_and_scheme
from django.contrib.auth.decorators import login_required, user_passes_test
from django.http import HttpResponseRedirect
from django.urls import reverse
from django.views.decorators.http import require_POST
import json
from django.core.mail import send_mail
from django.contrib.auth.models import Group
from django.core.paginator import Paginator
from django.http import HttpResponseForbidden
from .utils import *
from django.contrib.sessions.models import Session
from django.db.models import Sum, Count
from datetime import datetime
import datetime
from django.db.models import Q
from calendar import month_name
from datetime import datetime, timedelta
from django.utils.timezone import now
from collections import OrderedDict
import csv
from django.views.decorators.csrf import csrf_exempt
from .utils import is_admin, is_client
from django.views.decorators.http import require_GET
from django.utils import timezone
import logging
from collections import defaultdict
from decimal import Decimal, InvalidOperation
# Create your views here.
from django.core.mail import send_mail
from django.core.mail import EmailMultiAlternatives
from django.template.loader import render_to_string
from django.core.mail import EmailMessage
from dateutil.relativedelta import relativedelta
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def notify_payment_success(request, amount, reference):
    send_mail(
        subject="Payment Successful",
        message=(
            f"Hi {request.user.username},\n\n"
            f"Your payment of ₦{amount} was successful.\n"
            f"Transaction Reference: {reference}\n\n"
            f"Thank you for your trust!"
        ),
        from_email=settings.DEFAULT_FROM_EMAIL,
        recipient_list=[request.user.email],
        fail_silently=False  # Set to False for debugging
    )

def myHome(request):
    getSlider = SliderImages.objects.all()
    getAgent = Agent.objects.all()
    getPropertyType = PropertyType.objects.distinct()
    ClientTestimonies = ClientTestimony.objects.all()
    getPropertyCount = PropertyType.objects.annotate(property_count=Count('property'))
    getPropertyLocation = Location.objects.all().distinct()

    # Group by property_type (ForeignKey), get distinct types
    property_types = PropertyType.objects.all()

    grouped_properties = OrderedDict()
    for prop_type in property_types:
        grouped_properties[prop_type.name] = Property.objects.filter(
            property_type=prop_type
        ).prefetch_related(
            Prefetch('images', queryset=PropertyImage.objects.all())
        )

    context = {
        'getSliders': getSlider,
        'ClientTestimonies': ClientTestimonies,
        'getPropertyLocation': getPropertyLocation,
        'getPropertyType': getPropertyCount,
        'getPropertyType_': getPropertyType,
        'grouped_properties': grouped_properties,
        'property_types': property_types,
        'getAgent': getAgent
    }

    return render(request, 'index.html', context)

def about(request):
    return render(request, 'about.html')

def service(request):
    return render(request, 'services.html')

def contact(request):
    getAgent = Agent.objects.all()
    context= {'getAgent':getAgent}
    return render(request, 'contact.html', context)

@login_required(login_url='LoginUser')  # Ensure the login URL name is correct
def createSlider(request):
    if request.method == "POST":
        text = request.POST.get('sliderText')
        slider = request.FILES.get('sliderImage')
        if text and slider:
            SliderImages.objects.create(sliderText=text, sliderImage=slider)
            return JsonResponse({'status': 'success', 'message': 'Slider Image Uploaded Successfully'})
        else:
            return JsonResponse({'status': 'error', 'message': 'Missing data'})
    return render(request, 'createSlider.html')


@login_required(login_url='LoginUser')
def clientProfile(request):
    user = request.user
    is_client_user = is_client(user)
    is_admin_user = is_admin(user)

    if is_admin_user:
        # Annotate each user with the number of purchased products
        clientAll = User.objects.annotate(property_count=Count('purchasedproduct')).all()
        for user_obj in clientAll:
            user_obj.is_admin_user = is_admin(user_obj)
    
    elif is_client_user:
        clientAll = User.objects.annotate(property_count=Count('purchasedproduct')).filter(id=user.id)
        for user_obj in clientAll:
            user_obj.is_admin_user = False

    context = {
        'clients': clientAll,
        'is_client':is_client_user,
        'is_admin':is_admin_user
        }
    return render(request, 'Dashboard/ClientProfile.html', context)

@login_required(login_url='LoginUser')
def editProfile(request, user_id):
    user_to_edit = get_object_or_404(User, id=user_id)

    if request.method == 'POST':
        form = EditProfileForm(request.POST, instance=user_to_edit)
        if form.is_valid():
            form.save()
            messages.success(request, "Profile updated successfully.")
            return redirect('clientProfile')
    else:
        form = EditProfileForm(instance=user_to_edit)

    return render(request, 'Dashboard/EditProfile.html', {'form': form, 'user_to_edit': user_to_edit})
 


@login_required(login_url='LoginUser')
def propertyList(request):
    property_list = Property.objects.all().order_by('-date_added')
    paginator = Paginator(property_list, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    total_property_count = property_list.count()
    country_property_count = (
        Property.objects.values('location')
        .annotate(total=Count('id'))
        .order_by('-total')
    )

    # ✅ Create a mapping of property.id to list of clients
    property_clients_map = defaultdict(list)

    all_properties = Property.objects.all()
    all_purchases = PurchasedProduct.objects.select_related('user').all()

    # Match based on normalized title
    for prop in all_properties:
        matching_clients = [
            p.user for p in all_purchases
            if p.title.strip().lower() == prop.title.strip().lower()
        ]
        if matching_clients:
            property_clients_map[prop.id] = matching_clients

    transactions = Transaction.objects.filter(user=request.user).order_by('-created_at')[:10]

    context = {
        'page_obj': page_obj,
        'total_property_count': total_property_count,
        'country_property_count': country_property_count,
        'transactions': transactions,
        'property_clients_map': dict(property_clients_map),
    }

    return render(request, 'Dashboard/property_list.html', context)


@login_required(login_url='LoginUser')
def Dashboard(request):
    Google_API_KEY = settings.GOOGLE_MAPS_API_KEY
    user = request.user
    userFirstname = user.first_name
    is_client_user = is_client(user)
    is_admin_user = is_admin(user)

    # Initialize default values
    total_investment = 0
    to_balance = 0
    payment_progress = 0
    total_visitors = 0
    property_count = 0
    properties = Property.objects.none()  # Default empty queryset
    recent_payments = 0
    full_payment_history_url =0
    if is_admin_user:
        recent_payments = Transaction.objects.all().order_by('-created_at')[:5]
        # Admin sees all properties
        properties = Property.objects.all().order_by('-date_added')[:7]
        property_count = properties.count()
        purchased_xCount = PurchasedProduct.objects.all().count()
        # Sum deposits and outstanding balances across all purchases
        total_investment = PurchasedProduct.objects.aggregate(total=Sum('deposit'))['total'] or 0
        to_balance = PurchasedProduct.objects.aggregate(total=Sum('to_balance'))['total'] or 0

        # Count active sessions (as proxy for visitors)
        total_visitors = Session.objects.filter(expire_date__gte=datetime.now()).count()

    elif is_client_user:
        # Transaction History
        recent_payments = Transaction.objects.filter(user=user).order_by('-created_at')[:5]  # Fetch latest 5
        full_payment_history_url = reverse('payment_history')  # Define this URL in urls.py
        # Client sees only their purchased property titles
        purchased_titles = PurchasedProduct.objects.filter(user=user).values_list('title', flat=True)
        purchased_xCount = PurchasedProduct.objects.filter(user=user).count()
        # Filter Property based on titles
        properties = Property.objects.prefetch_related('images').filter(title__in=purchased_titles).order_by('-date_added')[:5]
        property_count = properties.count()

        # Calculate investment and balance for this user
        total_investment = PurchasedProduct.objects.filter(user=user).aggregate(total=Sum('deposit'))['total'] or 0
        to_balance = PurchasedProduct.objects.filter(user=user).aggregate(total=Sum('to_balance'))['total'] or 0
        total_price = PurchasedProduct.objects.filter(user=user).aggregate(total=Sum('price'))['total'] or 0

        # Calculate payment progress percentage
        if total_price and total_price > 0:
            payment_progress = round((total_investment / total_price) * 100, 2)

    # Fetch user's transactions (if any)
    transactions = Transaction.objects.filter(user=user).order_by('-created_at')
    getDate = datetime.now().year
    # Context passed to the template
    # Chart Data
    labels = []
    data = []

    today = now()
    for i in range(5, -1, -1):
        target_date = today - timedelta(days=30 * i)
        month_label = target_date.strftime('%b %Y')
        labels.append(month_label)

        if is_admin_user:
            total = Transaction.objects.filter(
                created_at__year=target_date.year,
                created_at__month=target_date.month
            ).aggregate(total=Sum('amount'))['total'] or 0
        else:
            total = Transaction.objects.filter(
                user=user,
                created_at__year=target_date.year,
                created_at__month=target_date.month
            ).aggregate(total=Sum('amount'))['total'] or 0

        data.append(round(total, 2))
    context = {
        'recent_payments': recent_payments,
        'full_payment_history_url': full_payment_history_url,
        'property': properties,
        'getDate': getDate,
        'purchased_xCount': purchased_xCount,
        'GoogleAPIKEY':Google_API_KEY,
        'propertyCount': property_count,
        'user': user,
        'userFirstname': userFirstname,
        'is_client': is_client_user,
        'is_admin': is_admin_user,
        'total_investment': total_investment,
        'to_balance': to_balance,
        'payment_progress': payment_progress,
        'total_visitors': total_visitors,
        'transactions': transactions,
        'chart_labels': labels,
        'chart_data': data,
    }

    return render(request, 'Dashboard/index.html', context)

# autocreate wallet
# @login_required(login_url="LoginUser")
# def wallet_view(request):
#     wallet = Wallet.objects.get(user=request.user)
#     return render(request, 'Dashboard/wallet.html', {'wallet': wallet})


@login_required
def initiate_wallet_payment(request):
    user = request.user    
    is_client_user = is_client(user)
    is_admin_user = is_admin(user)

    wallet = 0  # Default

    # Handle wallet context
    if is_admin_user:
        wallet = Wallet.objects.all().order_by('-updated_at')  # QuerySet for template loop
    elif is_client_user:
        wallet, _ = Wallet.objects.get_or_create(user=user)  # Safer than .get()

    callable_url = request.build_absolute_uri(reverse('verify_wallet_payment'))

    if request.method == 'POST':
        try:
            amount = int(request.POST.get('amount')) * 100  # Convert to kobo
            if amount < 10000:  # ₦100 minimum
                raise ValueError("Minimum payment is ₦100")

            reference = str(uuid.uuid4())
            email = user.email

            # Save transaction record
            WalletTransaction.objects.create(
                user=user,
                amount=amount / 100,  # store in naira
                reference=reference
            )

            headers = {
                "Authorization": f"Bearer {settings.PAYSTACK_SECRET_KEY}",
                "Content-Type": "application/json",
            }
            
            data = {
                "email": email,
                "amount": amount,
                "reference": reference,
                "callback_url": callable_url,
            }

            response = requests.post('https://api.paystack.co/transaction/initialize', json=data, headers=headers)
            res_data = response.json()

            if res_data.get('status'):
                return redirect(res_data['data']['authorization_url'])
            else:
                return render(request, 'error.html', {'message': 'Paystack initialization failed.'})

        except Exception as e:
            return render(request, 'error.html', {'message': str(e)})

    context = {
        'wallet': wallet,
        'is_admin_user': is_admin_user,
        'is_client_user': is_client_user
    }
    return render(request, 'Dashboard/wallet.html', context)


@login_required
@user_passes_test(is_admin)
def wallet_history(request):
    query = request.GET.get('q')
    users = User.objects.select_related('wallet').prefetch_related('wallettransaction_set')

    if query:
        users = users.filter(
            Q(username__icontains=query) | Q(email__icontains=query)
        )

    total_balance = Wallet.objects.aggregate(total=Sum('balance'))['total'] or 0

    context = {
        'users': users,
        'total_balance': total_balance
    }
    return render(request, 'Dashboard/wallet_history.html', context)



@login_required(login_url="LoginUser")
@csrf_exempt
def verify_wallet_payment(request):
    wallet = 0
    user = request.user    
    reference = request.GET.get('reference')

    if not reference:
        return render(request, 'error.html', {'message': 'Invalid reference'})

    headers = {
        "Authorization": f"Bearer {settings.PAYSTACK_SECRET_KEY}",
    }

    url = f"https://api.paystack.co/transaction/verify/{reference}"
    response = requests.get(url, headers=headers)
    res_data = response.json()

    if res_data['status'] and res_data['data']['status'] == 'success':
        try:
            trans = WalletTransaction.objects.get(reference=reference)
        except WalletTransaction.DoesNotExist:
            return render(request, 'Dashboard/error.html', {'message': 'Transaction not found'})

        if not trans.verified:
            trans.verified = True
            trans.save()

            # Credit wallet
            wallet, created = Wallet.objects.get_or_create(user=trans.user)
            wallet.balance += trans.amount
            wallet.save()

        return render(request, 'Dashboard/walletsuccess.html', {'amount': trans.amount})
    else:
        return render(request, 'Dashboard/wallet.html', {'message': 'Verification Failed'})

def error(request):
    return render(request,'Dashboard/error.html')

logger = logging.getLogger(__name__)
@login_required(login_url="LoginUser")
@csrf_exempt  
def verify_paystack_payment(request):
    reference = request.GET.get('reference')
    print("🔍 Verifying reference:", reference)
    if not reference:
        return JsonResponse({'status': False, 'message': 'Missing reference'})

    # Step 1: Call Paystack API
    headers = {
        "Authorization": f"Bearer {settings.PAYSTACK_SECRET_KEY}",
    }

    url = f"https://api.paystack.co/transaction/verify/{reference}"
    response = requests.get(url, headers=headers)
    res_data = response.json()

    if not res_data.get('status'):
        return JsonResponse({'status': False, 'message': 'Invalid Paystack response'})

    if res_data['data']['status'] != 'success':
        return JsonResponse({'status': False, 'message': 'Payment not successful'})

    # Step 2: Process successful payment
    try:
        transaction = WalletTransaction.objects.get(reference=reference)
    except WalletTransaction.DoesNotExist:
        return JsonResponse({'status': False, 'message': 'Transaction not found'})

    if transaction.verified:
        return JsonResponse({'status': True, 'message': 'Already verified'})

    # Mark transaction as verified
    transaction.verified = True
    transaction.save()

    # Optionally: update PurchasedProduct to reduce balance or mark as paid
    try:
        order = PurchasedProduct.objects.get(user=transaction.user, id__in=[
            int(ref) for ref in reference.split("-") if ref.isdigit()
        ])
        order.deposit += transaction.amount
        order.to_balance = max(order.price - order.deposit, 0)
        order.save()
    except PurchasedProduct.DoesNotExist:
        pass  # Optional: handle or log

    return JsonResponse({'status': True, 'message': 'Payment verified successfully'})

@login_required(login_url="LoginUser")
def payment_history(request):
    user = request.user
    is_client_user = is_client(user)
    is_admin_user = is_admin(user)

    status_filter = request.GET.get('status')
    reference_filter = request.GET.get('reference')

    # Filter the transactions
    queryset = Transaction.objects.all().order_by('-created_at') if is_admin_user else Transaction.objects.filter(user=user).order_by('-created_at')

    if status_filter:
        queryset = queryset.filter(status__iexact=status_filter)
    
    if reference_filter:
        queryset = queryset.filter(reference__iexact=reference_filter)

    # Fetch amount paid and to_balance
    if is_admin_user:
        total_paid = PurchasedProduct.objects.aggregate(total=Sum('deposit'))['total'] or 0
        total_balance = PurchasedProduct.objects.aggregate(total=Sum('to_balance'))['total'] or 0
    else:
        total_paid = PurchasedProduct.objects.filter(user=user).aggregate(total=Sum('deposit'))['total'] or 0
        total_balance = PurchasedProduct.objects.filter(user=user).aggregate(total=Sum('to_balance'))['total'] or 0

    for tx in queryset:
        products = tx.products.all()  # thanks to related_name='products'

        if products.exists():
            product = products.first()

            # Safely get deposit and balance values
            deposit = product.deposit if product.deposit is not None else 0
            balance = product.to_balance if product.to_balance is not None else 0

            tx.description = f"{product.title} | Deposit: ₦{deposit:,.2f} | Balance: ₦{balance:,.2f}"
        else:
            tx.description = "No product attached"


    # Payment progress
    payment_progress = 0
    total_price = total_paid + total_balance
    if total_price > 0:
        payment_progress = round((total_paid / total_price) * 100, 2)

    # CSV Export
    if request.GET.get('export') == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename=payment_history.csv'
        writer = csv.writer(response)
        writer.writerow(['Amount', 'User', 'Status', 'Reference', 'Date'])

        for t in queryset:
            writer.writerow([
                t.amount,
                t.user.username,
                t.status,
                t.reference or '-',
                t.created_at.strftime('%Y-%m-%d %H:%M')
            ])

        # Add summary
        writer.writerow([])
        writer.writerow(['Summary'])
        writer.writerow(['Total Paid', total_paid])
        writer.writerow(['Outstanding Balance', total_balance])
        writer.writerow(['Payment Progress (%)', f"{payment_progress}%"])

        return response

    context = {
        'all_payments': queryset,
        'payment_count': queryset.count(),
        'is_admin': is_admin_user,
        'is_client': is_client_user,
        'status_filter': status_filter,
        'total_paid': total_paid,
        'total_balance': total_balance,
        'payment_progress': payment_progress,
    }
    return render(request, 'Dashboard/payment_history.html', context)


@login_required
def delete_payment(request, payment_id):
    user = request.user    
    is_admin_user = is_admin(user)

    payment = get_object_or_404(Transaction, id=payment_id)

    if is_admin_user:
        # Backup the transaction
        PaymentTransactionTrash.objects.create(
            user=payment.user,
            amount=payment.amount,
            status=payment.status,
            reference=payment.reference,
            created_at=payment.created_at,
            is_trashed=True
        )

        # Backup and delete each related product
        purchased_products = PurchasedProduct.objects.filter(transaction=payment)
        for product in purchased_products:
            PurchasedProductTrash.objects.create(
                user=product.user,
                transaction_reference=payment.reference,
                title=product.title,
                method=product.method,
                price=product.price,
                deposit=product.deposit,
                to_balance=product.to_balance
            )
        purchased_products.delete()
        payment.delete()

        messages.success(request, "Payment and associated products moved to trash.")
    else:
        messages.error(request, "You are not authorized to delete this payment.")

    return redirect('payment_history')
# view trashed transaction
@login_required
def trashed_payments(request):
    user  = request.user
    is_client_user = is_client(user)
    is_admin_user = is_admin(user)
    if is_admin_user:
        trashed = PaymentTransactionTrash.objects.all().order_by('-created_at')
    else:
        trashed = PaymentTransactionTrash.objects.filter(user=request.user).order_by('-created_at')

    context = {
        'is_client': is_client_user,
        'is_admin': is_admin_user,
        'trashed_payments': trashed
    }
    return render(request, 'Dashboard/TransactionTrash.html', context)

@login_required(login_url='LoginUser')
def restore_payment(request, payment_id):
    # Get trashed transaction
    trash = get_object_or_404(PaymentTransactionTrash, id=payment_id)

    # Get all trashed products using the transaction reference
    trashed_products = PurchasedProductTrash.objects.filter(transaction_reference=trash.reference)

    # Recreate the Transaction
    restored_transaction = Transaction.objects.create(
        user=trash.user,
        amount=trash.amount,
        status=trash.status,
        reference=trash.reference,
        created_at=trash.created_at  # Note: might be overridden unless using raw SQL
    )

    # Recreate each PurchasedProduct
    for prod in trashed_products:
        PurchasedProduct.objects.create(
            user=prod.user,
            transaction=restored_transaction,
            title=prod.title,
            method=prod.method,
            price=prod.price,
            deposit=prod.deposit,
            to_balance=prod.to_balance
        )

    # Clean up trash records
    trash.delete()
    trashed_products.delete()

    messages.success(request, "Payment and products restored successfully.")
    return redirect('trashed_payments')



@login_required(login_url='LoginUser')
def permanent_delete_payment(request, payment_id):
    trash = get_object_or_404(PaymentTransactionTrash, id=payment_id)
    trash.delete()
    messages.success(request, "Payment permanently deleted.")
    return redirect('trashed_payments')

@login_required(login_url='LoginUser') 
def propertyList(request):
    # Get all properties
    user= request.user
    property_list = Property.objects.distinct().order_by('-date_added').prefetch_related('images')
    userFirstname =  user.first_name
    if not is_admin(user):
        return HttpResponseForbidden("Access Denied: Only admin can access this page.")
    # Pagination (6 per page)
    paginator = Paginator(property_list, 3)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Total number of all properties
    total_property_count = Property.objects.count()

    # Group by location (change to 'country' if needed)
    country_property_count = Property.objects.values('location').annotate(total=Count('id')).order_by('-total')

    # Transactions (optional)
    transactions = Transaction.objects.filter(user=request.user).order_by('-created_at')

    context = {
        'page_obj': page_obj,
        'total_property_count': total_property_count,
        'country_property_count': country_property_count,
        'transactions': transactions,
        'is_admin':True, 'user':user, 'userFirstname':userFirstname,
    }
    return render(request, 'Dashboard/property_list.html', context)

def FormProperty(request):
    ImageFormSet = modelformset_factory(PropertyImage, form=PropertyImageForm, extra=4, can_delete=False)

    if request.method == 'POST':
        form = PropertyForm(request.POST)
        formset = ImageFormSet(request.POST, request.FILES, queryset=PropertyImage.objects.none())

        if form.is_valid() and formset.is_valid():
            prop = form.save()
            for image_form in formset:
                if image_form.cleaned_data.get('image') and image_form.cleaned_data.get('label'):
                    PropertyImage.objects.create(
                        property=prop,
                        image=image_form.cleaned_data['image'],
                        label=image_form.cleaned_data['label']
                    )
            messages.success(request, "Property created successfully!")
            return redirect('FormProperty')  # Redirect back to the same form or another page
    else:
        form = PropertyForm()
        formset = ImageFormSet(queryset=PropertyImage.objects.none())

    return render(request, 'Dashboard/property.html', {
        'form': form,
        'image_forms': formset
    })

def propertyType(request):
    listPropertyType = propertyType.objects.unique()
    if request.method == 'POST':
        FormProperty = PropertyTypeForm(request.POST)

        if FormProperty.is_valid():
            FormProperty.save()
            messages.success(request, 'Property Type Inserted Successfully')
            return redirect('propertyType')
    else:
        FormProperty = PropertyTypeForm()
        getPropertyType =  PropertyType.objects.all()
    context = {'FormProperty': FormProperty, 'property-type':getPropertyType, 'listPropertyType':listPropertyType}
    return render(request, 'Dashboard/propertyType.html',context )


def property_list(request):
    properties = Property.objects.prefetch_related('images').all()
    return render(request, 'Dashboard/property_list.html', {'propertyView':properties})


# form Property Type

def FormPropertyType(request):
    formPropertyType = PropertyTypeForm(request.POST)
    if formPropertyType.is_valid():
        type_name = formPropertyType.cleaned_data['name'].strip().lower()

        # Check if it already exists (case-insensitive)
        checkProperty = PropertyType.objects.filter(name__iexact=type_name).exists()
        if checkProperty:
            messages.warning(request, f"Property Type '{type_name}' already exists.")
        else:
            formPropertyType.save()
            messages.success(request, 'Property Type Inserted Successfully')
            return redirect('/Dashboard/propertyType')
    return render(request, 'Dashboard/propertyType.html', {'property_type_form':formPropertyType})

# form Property Type

def FormPropertyLocation(request):
    formPropertyLocation = PropertyLocationForm(request.POST)
    if formPropertyLocation.is_valid():
        formPropertyLocation.save()
        messages.success(request, 'Property Location Inserted Successfully')
        return redirect('/Dashboard/propertyLocation')
    else:

        formPropertyLocation = PropertyLocationForm()
        
    return render(request, 'Dashboard/propertyLocation.html', {'property_location_form':formPropertyLocation})



def agent_list(request):
    agents = Agent.objects.all()
    return render(request, 'agents/agent_list.html', {'agents': agents})

def add_agent(request):
    if request.method == 'POST':
        form = AgentForm(request.POST, request.FILES)
        if form.is_valid():
            type_name = form.cleaned_data['name'].strip().lower()
            AgentProperty = Agent.objects.filter(name__iexact=type_name).exists()
            if AgentProperty:
                messages.warning(request, f"Agent Info '{type_name}' already exists.")
            else:
                form.save()
                messages.success(request, 'Agent added successfully.')
                
                return redirect('/Dashboard/add_agent')
    else:
        form = AgentForm()    
                
    return render(request, 'Dashboard/agent_form.html', {'form': form})


def clientTestimony(request):
    if request.method == 'POST':
        form = ClientTestimonyForm(request.POST, request.FILES)
        if form.is_valid():
            type_name = form.cleaned_data['name'].strip().lower()
            ClientTest = ClientTestimony.objects.filter(name__iexact=type_name).exists()
            if ClientTest:
                messages.warning(request, f"Client Testimony Info '{type_name}' already exists.")
            else:
                form.save()
                messages.success(request, 'Client Testimony added successfully.')
                
                return redirect('/Dashboard/client_testimony')
    else:
        form = ClientTestimonyForm()    
                
    return render(request, 'Dashboard/client_testimony.html', {'form': form})


def get_cordinate(address):
    api_key = settings.GOOGLE_MAPS_API_KEY
    url = f"https://maps.googleapis.com/maps/api/geocode/json?address={address}&key={api_key}"
    response =  requests.get(url)
    data =  response.json()
    if data['status']=='OK':
        location = data['results'][0]['geometry']['location']
        return location['lat'], location['lng']
    return None, None


def viewProperty(request, id):
    getPropertyView = get_object_or_404(Property, id=id)
    images =  getPropertyView.images.all()
    address = getPropertyView.location
    lat, lng = get_cordinate(address)
    subscriptionPlan =  getPropertyView.subscription_plans.all()

    context = {
        'property':getPropertyView,
        'images':images,
        'lat':lat,
        'lng':lng,
        'google_maps_api_key':settings.GOOGLE_MAPS_API_KEY,
        'subscriptionPlan':subscriptionPlan
    }
    return render(request, 'viewProperty.html', context)


# registration
def register_client(request):
    if request.method == 'POST':
        form = ClientRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(
                form.cleaned_data['password'])
            user.save()

            # add to client group
            client_group, created =  Group.objects.get_or_create(name='client')
            user.groups.add(client_group)
            # Save profile
            ClientProfile.objects.create(
                user=user,
                phone=form.cleaned_data['phone'],
                address=form.cleaned_data['address']
            )
            messages.success(request, 'Account created successfully.')
            return redirect('LoginUser')
    else:
        form = ClientRegistrationForm()
    return render(request, 'register_client.html', {'form': form})

def LoginUser(request):
    next_url = request.GET.get('next') or request.POST.get('next') or None
    
    if request.method == "POST":
        form = LoginForm(request=request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            if url_has_allowed_host_and_scheme(next_url, allowed_hosts={request.get_host()}):
                return redirect(next_url)
            return redirect("Dashboard")
        else:
            messages.warning(request, 'Authentication Failed... Username or Password Error')
    else:
        form = LoginForm()

    context = {'form': form, 'next': next_url or ''}
    return render(request, 'login.html', context)

    

# logout view
def custom_logout_view(request):
    logout(request)
    request.session.flush()
    messages.success(request, 'you have successfully logout')
    return redirect('LoginUser')


@login_required(login_url='LoginUser')  # Ensure the login URL name is correct
@require_POST
def add_to_cart(request, id):
    property_item = get_object_or_404(Property, id=id)
    cart = request.session.get('cart', {})

    if str(id) in cart:
        return JsonResponse({'message': 'This property is already in your cart.'}, status=200)

    if request.method == 'POST':
        try:
            body = json.loads(request.body)

            percentage = int(body.get('percentage', 100))
            deposit_raw = body.get('deposit', 0)
            selected_plan_id = body.get('subscription_plan_id')

            duration_months = 0
            plan = None
            monthly_breakdown_dict = {}
            increase_percent = 0
            increase_n_months = 0
            monthly_payment = 0

            if selected_plan_id:
                try:
                    plan = SubscriptionPropertyPlan.objects.get(id=selected_plan_id, property=property_item)
                except SubscriptionPropertyPlan.DoesNotExist:
                    plan = None
            else:
                plan = property_item.subscription_plans.first()

            if plan:
                duration_months = plan.duration_months or 0
                monthly_payment = plan.monthly_payment or 0
                monthly_breakdown_dict = plan.monthly_breakdown or {}
                increase_percent = plan.increase_percentage or 0
                increase_n_months = plan.increase_every_n_months or 0

            try:
                deposit = Decimal(str(deposit_raw))
            except (InvalidOperation, ValueError):
                return JsonResponse({'message': 'Invalid deposit amount.'}, status=400)

            # Convert breakdown dict to list for serialization
            breakdown_serializable = [
                {'month': month, 'amount': str(amount)} for month, amount in monthly_breakdown_dict.items()
            ]

            monthly_total = sum(Decimal(str(value)) for value in monthly_breakdown_dict.values())

            # Debug print
            print("Selected Plan:", plan)
            print("Monthly Payment:", monthly_payment)
            print("Serialized Breakdown:", breakdown_serializable)

            cart[str(id)] = {
                'title': property_item.title,
                'duration_months': duration_months,
                'monthly_payment': str(monthly_payment),  # Convert to string for session safety
                'monthly_breakdown': monthly_breakdown_dict,  # raw dict (ok if simple)
                'monthly_breakdown_schedule': breakdown_serializable,
                'monthly_breakdown_total': str(monthly_total),
                'increase_by_percent': str(increase_percent),
                'increase_every_n_months': increase_n_months,
                'price': str(property_item.price),
                'initial_deposit_percent': str(percentage),
                'initial_deposit_amount': str(deposit),
                'to_balance': str(property_item.price - deposit),
                'image': property_item.images.first().image.url if property_item.images.exists() else ''
            }

            request.session['cart'] = cart
            request.session.modified = True
            print("Cart updated:", json.dumps(cart, indent=2))

            return JsonResponse({'message': 'Property added to your cart.'}, status=201)

        except Exception as e:
            print("Cart Error:", str(e))
            return JsonResponse({'message': 'Error processing cart item.', 'error': str(e)}, status=400)

    return JsonResponse({'message': 'Invalid request method.'}, status=405)



@login_required(login_url='LoginUser')
def view_cart(request):
    user = request.user
    username = user.username
    getClient = get_object_or_404(ClientProfile, user=user)
    getAddress = getClient.address
    getFullname = user.get_full_name
    cart = request.session.get('cart', {})

    total_price = 0
    total_deposit = 0
    total_to_balance =0
    has_explicit_deposit = False
    monthly_payment = 0

    for item in cart.values():
    # Convert and sanitize values
        price = float(item.get('price', 0))
        deposit = float(item.get('initial_deposit_amount', 0)) if item.get('initial_deposit_amount') else price
        duration = int(item.get('duration_months', 0))
        monthly_payment = Decimal(item.get('monthly_payment', 0))

        # Calculate outstanding balance
        to_balance = price - deposit

        # Update item with computed values
        item['to_balance'] = to_balance
        item['duration_months'] = duration  # Ensure it's int for consistent rendering

        # Update totals
        total_price += price
        total_deposit += deposit
        total_to_balance += to_balance  
        if isinstance(item.get('initial_deposit_percent'), str):
            item['initial_deposit_percent'] = int(item['initial_deposit_percent'])
        # Check if the user explicitly selected a deposit option
        if item.get('initial_deposit_amount'):
            has_explicit_deposit = True
        if item.get('duration_months'):
            item['duration_months'] = duration
    
    print(monthly_payment)
    context = {
        'username':username,
        'getAddress':getAddress,
        'getFullname': getFullname,
        'cart': cart,
        'total_to_balance':total_to_balance,
        'total_price': total_price,
        'total_deposit': total_deposit if has_explicit_deposit else None,
        'PAYSTACK_PUBLIC_KEY': settings.PAYSTACK_PUBLIC_KEY,
    }

    return render(request, 'cart.html', context)


@login_required
def update_balances_view(request):
    today = now()
    updated_count = 0
    updated_items = []

    products = PurchasedProduct.objects.select_related('property', 'user')
    for product in products:
        prop = product.property
        if not prop:
            continue

        plan = prop.subscription_plans.last()
        if not plan or not plan.increase_every_n_months or not plan.increase_percentage:
            continue

        months_since_purchase = (today.year - product.date_added.year) * 12 + (today.month - product.date_added.month)
        increments_due = months_since_purchase // plan.increase_every_n_months
        if increments_due <= 0:
            continue

        base_balance = product.price - (product.deposit or 0)
        new_balance = Decimal(str(base_balance))

        for _ in range(increments_due):
            new_balance += (Decimal(plan.increase_percentage) / 100) * new_balance

        new_balance = round(new_balance, 2)

        if product.to_balance != new_balance:
            old_balance = product.to_balance
            product.to_balance = new_balance
            product.save()
            updated_count += 1
            updated_items.append({
                "user": product.user.get_full_name() or product.user.username,
                "email": product.user.email,
                "property": prop.title,
                "old_balance": old_balance,
                "new_balance": new_balance,
                "increase": plan.increase_percentage,
            })

    if updated_items:
        # === Create Word Document ===
        doc = Document()

        # Insert Logo
        logo_path = 'static/images/logo.png'  # Update path if needed
        try:
            doc.add_picture(logo_path, width=Inches(1.5))
        except:
            pass  # Skip if logo not found

        # Title
        title = doc.add_paragraph("Property Balance Update Report")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(18)
        title.runs[0].bold = True

        doc.add_paragraph(f"📅 Date: {today.strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"✅ Total Updated: {updated_count}\n")

        # Add a border-like section for each update
        for entry in updated_items:
            p = doc.add_paragraph()
            run = p.add_run(f"User: {entry['user']} ({entry['email']})\n")
            run.bold = True
            p.add_run(f"Property: {entry['property']}\n")
            p.add_run(f"Old Balance: ₦{entry['old_balance']}\n")
            p.add_run(f"New Balance: ₦{entry['new_balance']}\n")
            p.add_run(f"Increase: {entry['increase']}%\n")
            doc.add_paragraph("—" * 35)  # Divider

        # Footer
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("This report was auto-generated by Wiltshire Property System.")
        footer_run.italic = True
        footer_run.font.size = Pt(9)

        # Save document to in-memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Send email with attachment
        email = EmailMessage(
            subject="📄 Property Balance Update - Withshire Developments",
            body="Please find attached the latest balance update report in Word format.",
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=["wonderpaul242@gmail.com", request.user.email],
        )
        email.attach('PropertyBalanceUpdate.docx', buffer.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        email.send()

    messages.success(request, f"{updated_count} purchased products were updated.")
    return redirect('Dashboard')


@login_required(login_url='LoginUser')
def remove_from_cart(request, id):
    cart = request.session.get('cart', {})
    if id in cart:
        del cart[id]
        request.session['cart'] = cart
        request.session.modified = True
    return HttpResponseRedirect(reverse('view_cart'))


def add_subscription_plan(request, property_id):
    property = get_object_or_404(Property, id=property_id)

    if property.allSubscription != 'Yes':
        messages.error(request, "This property does not support subscription plans.")
        return redirect('add_subscription_plan', property_id=property.id)

    # Try to fetch an existing plan (assuming one plan per property)
    existing_plan = SubscriptionPropertyPlan.objects.filter(property=property).first()

    if request.method == 'POST':
        form = SubscriptionPropertyPlanForm(request.POST, instance=existing_plan)
        if form.is_valid():
            plan = form.save(commit=False)
            plan.property = property

            total_price = property.price
            deposit_percent = plan.initial_deposit_percent

            # Calculate initial payment
            initial_payment = total_price * Decimal(deposit_percent / 100)
            plan.auto_balance = True
            plan.initial_payment = initial_payment

            # Calculate monthly payment
            if plan.duration_months > 1:
                remaining = total_price - initial_payment
                months_remaining = plan.duration_months
                plan.monthly_payment = remaining / Decimal(months_remaining)
            else:
                plan.monthly_payment = Decimal("0.00")

            plan.save()
            messages.success(request, "Subscription plan saved successfully.")
            return redirect('add_subscription_plan', property_id=property.id)
    else:
        form = SubscriptionPropertyPlanForm(instance=existing_plan)

    return render(request, 'Dashboard/add_plan.html', {
        'form': form,
        'property': property,
        'is_update': existing_plan is not None
    })

@login_required(login_url='LoginUser')
def payment_success(request):
    reference = request.GET.get('reference')
    # Optionally verify with Paystack API here using secret key
    messages.success(request, "Payment successful! Reference: ", reference)
    return render(request, 'some_thank_you_page_or_dashboard.html')



@login_required(login_url='LoginUser')
def verify_payment(request):
    reference = request.GET.get('reference')

    if not reference:
        messages.error(request, "No reference supplied.")
        return redirect('view_cart')

    # Avoid duplicate processing
    if Transaction.objects.filter(reference=reference).exists():
        messages.success(request, "Transaction already verified.")
        return redirect('paymentVerification', reference=reference)

    # Verify with Paystack
    headers = {
        "Authorization": f"Bearer {settings.PAYSTACK_SECRET_KEY}",
        "Content-Type": "application/json",
    }
    url = f"https://api.paystack.co/transaction/verify/{reference}"
    response = requests.get(url, headers=headers)
    res_data = response.json()

    contract_instance = UploadedContract.objects.filter(reference=reference).first()  # if saved in separate model

    if res_data.get('status') and res_data['data']['status'] == 'success':
        amount = res_data['data']['amount'] / 100  # Convert kobo to Naira

        # ✅ Create transaction
        transaction = Transaction.objects.create(
            user=request.user,
            reference=reference,
            amount=amount,
            contract_file=contract_instance.contract_file if contract_instance else None,  # 🧾 attach PDF
            status='success'
        )

        # ✅ Retrieve contract file uploaded earlier
        

        # ✅ Save purchased items
        cart_items = request.session.get('cart', {})
        for item in cart_items.values():
            price = float(item['price'])
            deposit = float(item.get('initial_deposit_amount', price))
            paid_percent = int(item.get('initial_deposit_percent', 100))
            to_balance = price - deposit
            property_instance = Property.objects.get(title=item['title'])
            product = PurchasedProduct.objects.create(
                user=request.user,
                transaction=transaction,
                property = property_instance,
                title= property_instance.title,
                method="Paystack",
                price=price,
                deposit=deposit,
                Paidpercentage=paid_percent,
                to_balance=to_balance,
                contract_file=contract_instance.contract_file if contract_instance else None  # 🧾 attach PDF
            )

        #  Optionally send email
        #  notify_payment_success(request, amount, reference)
        
        user =  request.user
        contract_file = contract_instance.contract_file  # FileField

        message = (
            f"Dear {user.get_full_name() or user.username},\n\n"
            f"Your payment of ₦{amount:,.2f} was successful.\n"
            f"Transaction Reference: {reference}\n\n"
            f"Attached is your property contract for Wiltshire Heights.\n"
            f"Thank you for doing business with us.\n\n"
            f"Wiltshire Estates Team"
        )

        email = EmailMessage(
            subject="Your Property Contract and Payment Confirmation",
            body=message,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=["wonderpaul242@gmail.com", user.email],
        )

        # ✅ Attach the contract file if it exists
        if contract_file:
            email.attach('Contract_Wiltshire.pdf', contract_file.read(), 'application/pdf')

        email.send()
        

        # Clear session cart
        request.session['cart'] = {}
        request.session.modified = True

        messages.success(request, f"Payment successful! Ref: {reference}")
        return redirect('paymentVerification', reference=reference)

    else:
        messages.error(request, "Payment verification failed.")
        return redirect('view_cart')


@csrf_exempt
def save_purchased_product(request):
    if request.method == 'POST':
        contract_file = request.FILES.get('contract_pdf')
        reference = request.POST.get('reference')

        if not contract_file or not reference:
            return JsonResponse({'status': 'error', 'message': 'Missing file or reference'}, status=400)

        UploadedContract.objects.create(
            reference=reference,
            contract_file=contract_file
        )
        return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'error', 'message': 'Invalid request'}, status=405)


@login_required(login_url='LoginUser')
def paymentVerification(request, reference):
    transaction = get_object_or_404(Transaction, reference=reference)
    products = transaction.products.all()  # using related_name from the model

    context = {
        'transaction': transaction,
        'products': products,
    }
    return render(request, 'Dashboard/payementVerification.html', context)



def compute_payment_details(price, deposit_percent, duration_months):
    """
    Calculate initial and monthly payment for a property subscription plan.
    """
    total_price = Decimal(price)
    deposit = total_price * Decimal(deposit_percent) / Decimal(100)
    remaining = total_price - deposit

    if duration_months > 1:
        monthly_payment = remaining / Decimal(duration_months)
    else:
        monthly_payment = Decimal("0.00")

    return round(deposit, 2), round(monthly_payment, 2)




@login_required(login_url="LoginUser")
def my_orders(request):
    user = request.user
    username = user.username

    # Check user roles
    is_client_user = is_client(user)
    is_admin_user = is_admin(user)

    # Get user orders
    if is_client_user:
        orders = PurchasedProduct.objects.filter(user=user)
    elif is_admin_user:
        orders = PurchasedProduct.objects.all()
    else:
        orders = PurchasedProduct.objects.none()

    # Admin filter by user
    selected_user_id = request.GET.get("user")
    if selected_user_id:
        orders = orders.filter(user_id=selected_user_id)

    # Filter by date
    filter_type = request.GET.get("filter")
    today = datetime.today()
    if filter_type == "day":
        orders = orders.filter(date_added__date=today.date())
    elif filter_type == "month":
        orders = orders.filter(date_added__month=today.month)
    elif filter_type == "year":
        orders = orders.filter(date_added__year=today.year)

    # Search
    search_query = request.GET.get("search")
    if search_query:
        q = search_query.lower()
        if q in ["completed", "fully paid", "paid", "done"]:
            orders = orders.filter(to_balance__lte=0)
        elif q in ["owing", "pending", "not paid", "balance", "with balance"]:
            orders = orders.filter(to_balance__gt=0)
        else:
            orders = orders.filter(
                Q(title__icontains=search_query) |
                Q(method__icontains=search_query)
            )

    orders = orders.order_by("-date_added")

    for order in orders:
        # Get subscription plan
        try:
            property_obj = Property.objects.get(title=order.title)
            subscription_plan = SubscriptionPropertyPlan.objects.filter(property=property_obj).first()
            order.subscription_plan = subscription_plan
        except Property.DoesNotExist:
            order.subscription_plan = None

        # Calculate payment amount
        amount = order.to_balance or 0

        # Try to reuse existing unverified transaction
        existing_txn = WalletTransaction.objects.filter(
            user=order.user,
            amount=amount,
            verified=False,
            reference__startswith=f"ORDER-{order.id}-"
        ).first()

        if existing_txn:
            order.payment_reference = existing_txn.reference
            
        else:
            fresh_reference = f"ORDER-{order.id}-{uuid4().hex[:8]}"
            WalletTransaction.objects.create(
                user=order.user,
                amount=amount,
                reference=fresh_reference,
                verified=False,
            )
            print(fresh_reference)
            order.payment_reference = fresh_reference
            

    # Percent options
    percentage_options = [str(i) for i in range(30, 101, 10)]

    # For admin filter dropdown
    admin_users = Group.objects.get(name="admin").user_set.all()
    non_admin_users = get_user_model().objects.exclude(id__in=admin_users)

    return render(request, "Dashboard/my_orders.html", {
        "orders": orders,
        "is_client": is_client_user,
        "is_admin": is_admin_user,
        "user_username": username,
        "percentage_options": percentage_options,
        "PayStackKey": settings.PAYSTACK_PUBLIC_KEY,
        "all_users": non_admin_users if request.user.is_superuser else None,
        "selected_user_id": selected_user_id,
    })


@login_required(login_url="LoginUser")
def process_payment(request, product_id):
    if request.method != 'POST':
        messages.error(request, "Invalid request method.")
        return redirect("my_orders")

    product = get_object_or_404(PurchasedProduct, id=product_id, user=request.user)
    method = request.POST.get("method")
    payment_type = request.POST.get("payment_option")  # corrected name
    balance = Decimal(product.to_balance or 0)

    if balance <= 0:
        messages.error(request, "No remaining balance to pay.")
        return redirect("my_orders")

    # Determine amount to pay
    if payment_type == "percentage":
        try:
            percent = float(request.POST.get("percentage", 0))
            if percent < 30 or percent > 100:
                messages.error(request, "Percentage must be between 30% and 100%.")
                return redirect("my_orders")
            amount = (Decimal(percent) / 100) * balance
        except ValueError:
            messages.error(request, "Invalid percentage value.")
            return redirect("my_orders")
    else:
        amount = balance / 6  # Monthly payment

    # Wallet payment
    if method == "wallet":
        wallet = getattr(request.user, 'wallet', None)

        if not wallet:
            messages.error(request, "Wallet not found for user.")
            return redirect("my_orders")

        if wallet.balance < amount:
            messages.error(request, "Insufficient wallet balance.")
            return redirect("my_orders")

        wallet.balance -= amount
        wallet.save()

        Transaction.objects.create(
            user=request.user,
            reference=str(uuid.uuid4()),
            amount=amount,
            status="success"
        )

        product.to_balance = float(balance - amount)
        product.save()

        messages.success(request, f"Wallet payment of ₦{amount:.2f} successful.")
        return redirect("my_orders")
    
    # Paystack payment
    elif method == "paystack":
        paystack_secret = settings.PAYSTACK_SECRET_KEY
        callable_url = request.build_absolute_uri(reverse('verify_paystack_payment'))

        headers = {
            "Authorization": f"Bearer {paystack_secret}",
            "Content-Type": "application/json",
        }

        data = {
            "email": request.user.email,
            "amount": int(amount * 100),  # in kobo
            "reference": str(uuid.uuid4()),
            "callback_url": callable_url,
        }

        response = requests.post("https://api.paystack.co/transaction/initialize", json=data, headers=headers)
        res_data = response.json()

        if response.status_code == 200 and res_data['status']:
            # Save pending transaction
            Transaction.objects.create(
                user=request.user,
                reference=data["reference"],
                amount=amount,
                status="pending"
            )
            return redirect(res_data['data']['authorization_url'])
        else:
            messages.error(request, "Failed to initiate Paystack payment.")
            return redirect("my_orders")

    else:
        messages.error(request, "Unsupported payment method.")
        return redirect("my_orders")

def property_lists(request):
    query = request.GET.get("q", "")
    plan = request.GET.get("plan", "")
    location = request.GET.get("location", "")
    min_price = request.GET.get("min_price", "")
    max_price = request.GET.get("max_price", "")

    properties = Property.objects.all()

    if query:
        properties = properties.filter(
            Q(title__icontains=query) |
            Q(location__city__icontains=query) |
            Q(location__state__icontains=query) |
            Q(location__address__icontains=query)
        )


    if plan == "subscription":
        properties = properties.filter(allSubscription='Yes')

    if location:
        properties = properties.filter(location__id=location)

    if min_price:
        properties = properties.filter(price__gte=min_price)

    if max_price:
        properties = properties.filter(price__lte=max_price)

    paginator = Paginator(properties, 15)
    page = request.GET.get("page")
    properties = paginator.get_page(page)

    all_locations = Location.objects.order_by('state').distinct()

    return render(request, "properties.html", {
        "properties": properties,
        "all_locations": all_locations,
        "now": timezone.now(),
    })

def PropertyTypeHome(request):
    getSlider = SliderImages.objects.all()
    getAgent = Agent.objects.all()
    ClientTestimonies = ClientTestimony.objects.all()
    getPropertyLocation = Location.objects.all().distinct()

    # Property type list & count
    property_types = PropertyType.objects.all()
    getPropertyCount = property_types.annotate(property_count=Count('property'))

    # Group properties by type with related images
    grouped_properties = OrderedDict()
    for prop_type in property_types:
        properties = Property.objects.filter(
            property_type=prop_type
        ).prefetch_related(
            Prefetch('images', queryset=PropertyImage.objects.all())
        )
        grouped_properties[prop_type.name] = properties

    context = {
        'getSliders': getSlider,
        'ClientTestimonies': ClientTestimonies,
        'getPropertyLocation': getPropertyLocation,
        'getPropertyType': getPropertyCount,
        'getPropertyType_': property_types,
        'grouped_properties': grouped_properties,
        'property_types': property_types,
        'getAgent': getAgent
    }

    return render(request, 'PropertyTypeHome.html', context)


# views.py

@login_required
@require_GET
def pay_with_wallet(request):
    try:
        cart = request.session.get('cart', {})
        if not cart:
            return JsonResponse({'status': 'fail', 'message': 'Your cart is empty.'}, status=400)

        wallet = Wallet.objects.select_for_update().get(user=request.user)
        total_deposit = Decimal('0.00')

        with transaction.atomic():
            for item in cart.values():
                deposit_raw = item.get('initial_deposit_amount') or item.get('price')
                value = safe_decimal(deposit_raw)

                if value is None:
                    return JsonResponse({'status': 'fail', 'message': 'Invalid item amount in cart.'}, status=400)

                total_deposit += value

            if wallet.balance < total_deposit:
                return JsonResponse({'status': 'fail', 'message': 'Insufficient wallet balance.'}, status=400)

            # Deduct from wallet
            wallet.balance -= total_deposit
            wallet.save()

            # Create a Transaction record
            reference = f"WL-{uuid.uuid4().hex[:12]}"
            transaction_record = Transaction.objects.create(
                user=request.user,
                amount=total_deposit,
                reference=reference,
                status='success'
            )

            # Create PurchasedProduct entries
            for property_id, item in cart.items():
                property_obj = Property.objects.filter(id=property_id).first()
                price = safe_decimal(item.get('price'))
                deposit = safe_decimal(item.get('initial_deposit_amount') or price)
                to_balance = (price or Decimal('0.00')) - (deposit or Decimal('0.00'))

                PurchasedProduct.objects.create(
                    user=request.user,
                    transaction=transaction_record,                    
                    title=item.get('title', 'Unknown'),
                    price=price,
                    method="Wallet",
                    deposit=deposit,
                    to_balance=to_balance
                )

            # Clear cart
            request.session['cart'] = {}
            request.session.modified = True

        return JsonResponse({
            'status': 'success',
            'message': 'Payment successful using wallet.',
            'redirect_url': reverse('Dashboard')
        })

    except Wallet.DoesNotExist:
        return JsonResponse({'status': 'fail', 'message': 'Wallet not found.'}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'fail', 'message': f'Server error: {str(e)}'}, status=500)
    

def is_admin(user):
    return user.is_staff or user.is_superuser

@login_required
@user_passes_test(is_admin)
def chat_dashboard(request):
    clients = User.objects.exclude(id=request.user.id)
    return render(request, 'Dashboard/chat.html', {'clients': clients})


@login_required
@user_passes_test(is_admin)
def chat_with_user(request, user_id):
    user =  request.user
    is_admin_user = is_admin(user)
    is_client_user = is_client(user)
    recipient = get_object_or_404(User, id=user_id)
    messages = Message.objects.filter(
        sender__in=[request.user, recipient],
        recipient__in=[request.user, recipient]
    ).order_by('timestamp')

    if request.method == 'POST':
        content = request.POST.get('content')
        if content:
            Message.objects.create(sender=request.user, recipient=recipient, content=content)
            return redirect('chat_with_user', user_id=recipient.id)

    clients = User.objects.exclude(id=request.user.id)
    return render(request, 'Dashboard/chat.html', {
        'clients': clients,
        'recipient': recipient,
        'messages': messages,
        'is_client':is_client_user,
        'is_admin':is_admin_user,
    })


@login_required
def client_chat_view(request):
    user = request.user
    is_admin_user = is_admin(user)
    is_client_user = is_client(user)
    admin_user = User.objects.filter(is_superuser=True).first()
    if not admin_user:
        return render(request, 'Dashboard/chat_client.html', {'error': 'No admin available at the moment.'})

    messages = Message.objects.filter(
        Q(sender=request.user, recipient=admin_user) |
        Q(sender=admin_user, recipient=request.user)
    ).order_by('timestamp')

    if request.method == 'POST':
        content = request.POST.get('content')
        if content:
            Message.objects.create(sender=request.user, recipient=admin_user, content=content)
            return redirect('client_chat')

    return render(request, 'Dashboard/chat_client.html', {
        'admin_user': admin_user,
        'messages': messages,
        'is_client':is_client_user,
        'is_admin':is_admin_user,
    })


# @csrf_exempt
# def toggle_autobalance(request, client_id):
#     if request.method == "POST":
#         data = json.loads(request.body)
#         status = data.get("status")

#         client = Client.objects.get(id=client_id)
#         client.autobalance = status
#         client.save()

#         return JsonResponse({"success": True, "autobalance": client.autobalance})
#     return JsonResponse({"error": "Invalid request"}, status=400)